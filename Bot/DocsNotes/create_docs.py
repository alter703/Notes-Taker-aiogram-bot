from io import BytesIO
from docx import Document

from Bot.DatabaseQueries.getters import get_one_note


async def create_dock(user_id, title):
    row = await get_one_note(user_id=user_id, title=title)

    note_as_dock = Document()

    note_as_dock.add_heading(f"{row[0].get('row')[0]}", 2)

    paragraph = note_as_dock.add_paragraph("")
    paragraph.add_run(f"{row[0].get('row')[1]}")

    buf = BytesIO()

    note_as_dock.save(buf)
    buf.seek(0)

    return buf
